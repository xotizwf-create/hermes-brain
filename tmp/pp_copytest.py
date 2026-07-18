"""E2E: copy_contract + edit_contract (monthly schedule pattern) on seeded data (miramed32)."""
import io
import json
import re
import time
import urllib.request
from pathlib import Path

APP = Path("/var/www/prostye-postavki/app")
import sys
sys.path.insert(0, str(APP))
PENDING = APP / "backend" / "data" / "pending_contract_files"
DOC_ID = "smoketest-schedule"

env = {}
for line in (APP / ".env.local").read_text(encoding="utf-8", errors="replace").splitlines():
    if "=" in line and not line.startswith("#"):
        k, _, v = line.partition("=")
        env[k.strip()] = v.strip().strip('"')
URL = f"http://127.0.0.1:8000/mcp/{env['MCP_SERVER_SECRET']}"


def call(name, arguments):
    body = json.dumps({"jsonrpc": "2.0", "id": 1, "method": "tools/call",
                       "params": {"name": name, "arguments": arguments}}).encode()
    req = urllib.request.Request(URL, data=body, headers={"content-type": "application/json"})
    with urllib.request.urlopen(req, timeout=180) as resp:
        p = json.loads(resp.read().decode())
    if "error" in p:
        raise RuntimeError(f"{name}: {p['error']}")
    return json.loads(p["result"]["content"][0]["text"])


rows = call("read_table_rows", {"table": "organizations", "limit": 200})
orgs = []
for r in rows.get("rows", []):
    inn = re.sub(r"\D", "", str(r.get("inn") or ""))
    if len(inn) in (10, 12) and not str(r.get("name", "")).startswith("__SYSTEM_") and r.get("is_active") in (True, None):
        orgs.append(inn)
customer_inn, supplier_inn = orgs[0], orgs[1]

from docx import Document
doc = Document()
doc.add_paragraph("КОНТРАКТ ПОСТАВКИ № ГРАФИК-ТЕСТ/1")
doc.add_paragraph(f"«01» июля 2026 г. Заказчик ИНН {customer_inn}, Поставщик ИНН {supplier_inn}.")
doc.add_paragraph("График: июль 10 шт, август 20 шт, сентябрь 30 шт по 100,00 руб.")
buf = io.BytesIO()
doc.save(buf)
content = buf.getvalue()
PENDING.mkdir(parents=True, exist_ok=True)
(PENDING / f"{DOC_ID}.bin").write_bytes(content)
(PENDING / f"{DOC_ID}.json").write_text(json.dumps({
    "id": DOC_ID, "filename": "ТЕСТ график — УДАЛИТЬ.docx", "size_bytes": len(content),
    "uploaded_at": int(time.time()),
    "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}, ensure_ascii=False), encoding="utf-8")

ids = []
try:
    # период 1 (июль) — из документа
    res1 = call("save_contract_from_incoming_document", {
        "documentId": DOC_ID,
        "extracted": {
            "contractNumber": "ГРАФИК-ТЕСТ/1", "contractDate": "01.07.2026", "deadlineDate": "31.07.2026",
            "customerInn": customer_inn, "supplierInn": supplier_inn, "vatPercent": "0",
            "items": [{"name": "Тестовый товар графика", "docName": "Тестовый товар графика",
                       "unit": "шт", "qty": 10, "price": 100}],
        },
    })
    ids.append(str(res1["contract"]["id"]))
    print(f"период 1 создан: дедлайн {res1['deadlineSaved']}, total {res1['totalAmount']}")

    # период 2 (август) — copy_contract с itemUpdates
    res2 = call("copy_contract", {
        "contractId": ids[0], "deadlineDate": "31.08.2026",
        "itemUpdates": [{"name": "Тестовый товар графика", "qty": 20}],
    })
    ids.append(res2["contract"]["contractId"])
    c2 = res2["contract"]
    assert c2["number"] == "Контракт ГРАФИК-ТЕСТ/1" or "ГРАФИК-ТЕСТ/1" in c2["number"], c2["number"]
    assert c2["deadline"] == "31.08.2026", c2["deadline"]
    assert c2["items"][0]["qty"] == 20 and c2["items"][0]["price"] == 100, c2["items"]
    assert c2["totalAmount"] == 2000.0 and c2["status"] == "active", c2
    files2 = call("get_contract_files", {"contractId": ids[1]})
    assert files2["count"] >= 1, "файл не склонировался в копию"
    print(f"период 2 (copy): дедлайн {c2['deadline']}, qty 20, total {c2['totalAmount']}, файл склонирован OK")

    # период 3 (сентябрь) — copy с полной заменой items
    res3 = call("copy_contract", {
        "contractId": ids[0], "deadlineDate": "30.09.2026",
        "items": [{"name": "Тестовый товар графика", "unit": "шт", "qty": 30, "price": 100}],
    })
    ids.append(res3["contract"]["contractId"])
    c3 = res3["contract"]
    assert c3["deadline"] == "30.09.2026" and c3["totalAmount"] == 3000.0, c3
    print(f"период 3 (copy, items replace): дедлайн {c3['deadline']}, total {c3['totalAmount']} OK")

    # edit_contract: пометить период 1 завершённым + факт
    res_e = call("edit_contract", {
        "contractId": ids[0], "status": "completed", "factDate": "15.07.2026",
        "itemUpdates": [{"index": 1, "factQty": 10, "factDate": "15.07.2026"}],
    })
    ce = res_e["contract"]
    assert ce["status"] == "completed", ce
    print(f"edit: период 1 → completed, изменения: {res_e['changes']}")

    # edit: точечная правка цены
    res_p = call("edit_contract", {"contractId": ids[1], "itemUpdates": [{"index": 1, "price": 110}]})
    assert res_p["contract"]["totalAmount"] == 2200.0, res_p["contract"]
    print(f"edit: цена периода 2 → 110, total {res_p['contract']['totalAmount']} OK")

    # общая картина: 3 записи одного номера с разными дедлайнами
    got = call("get_contracts", {"query": "ГРАФИК-ТЕСТ", "limit": 10})
    cnt = got.get("count") or len(got.get("contracts") or got.get("rows") or [])
    print(f"записей по номеру: {cnt}")
    print("SCHEDULE_TOOLS_OK")
finally:
    import psycopg
    db_url = (env.get("DATABASE_URL") or "").replace("postgresql+psycopg2://", "postgresql://")
    with psycopg.connect(db_url) as conn:
        with conn.cursor() as cur:
            for cid in ids:
                cur.execute("SELECT storage_key FROM contract_files WHERE contract_id = %s", (cid,))
                for (key,) in cur.fetchall():
                    p = (APP / key).resolve()
                    if "contract_files_store" in str(p) and p.exists():
                        p.unlink()
                cur.execute("DELETE FROM inventory_movements WHERE contract_id = %s", (cid,))
                cur.execute("DELETE FROM delivery_stage_items WHERE stage_id IN (SELECT id FROM delivery_stages WHERE contract_id = %s)", (cid,))
                cur.execute("DELETE FROM delivery_stages WHERE contract_id = %s", (cid,))
                cur.execute("DELETE FROM contract_items WHERE contract_id = %s", (cid,))
                cur.execute("DELETE FROM contract_files WHERE contract_id = %s", (cid,))
                cur.execute("DELETE FROM contracts WHERE id = %s", (cid,))
            cur.execute("DELETE FROM pending_document_ai_state WHERE doc_id = %s", (DOC_ID,))
        conn.commit()
    for suffix in (".json", ".bin", ".preview.bin"):
        p = PENDING / f"{DOC_ID}{suffix}"
        if p.exists():
            p.unlink()
    print("cleanup done")
