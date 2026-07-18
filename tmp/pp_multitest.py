"""E2E: multi-contract document via keepDocument (runs on miramed32).

Seeds one DOCX with TWO contracts, creates contract A with keepDocument=true
(doc must stay in incoming, file attached to A), then contract B without it
(doc consumed, file attached to B). Full cleanup. Real docs untouched.
"""
import io
import json
import re
import time
import urllib.request
from pathlib import Path

APP = Path("/var/www/prostye-postavki/app")
import sys
sys.path.insert(0, str(APP))
PENDING = APP / "backend" / "data" / "pending_contract_files"
DOC_ID = "smoketest-multi"

env = {}
for line in (APP / ".env.local").read_text(encoding="utf-8", errors="replace").splitlines():
    if "=" in line and not line.startswith("#"):
        k, _, v = line.partition("=")
        env[k.strip()] = v.strip().strip('"')
URL = f"http://127.0.0.1:8000/mcp/{env['MCP_SERVER_SECRET']}"


def call(name, arguments):
    body = json.dumps({"jsonrpc": "2.0", "id": 1, "method": "tools/call",
                       "params": {"name": name, "arguments": arguments}}).encode()
    req = urllib.request.Request(URL, data=body, headers={"content-type": "application/json"})
    with urllib.request.urlopen(req, timeout=180) as resp:
        p = json.loads(resp.read().decode())
    if "error" in p:
        raise RuntimeError(f"{name}: {p['error']}")
    return json.loads(p["result"]["content"][0]["text"])


# real org INNs from the directory (no new orgs, no DaData)
rows = call("read_table_rows", {"table": "organizations", "limit": 200})
orgs = []
for r in rows.get("rows", []):
    inn = re.sub(r"\D", "", str(r.get("inn") or ""))
    if len(inn) in (10, 12) and not str(r.get("name", "")).startswith("__SYSTEM_") and r.get("is_active") in (True, None):
        orgs.append(inn)
customer_inn, supplier_inn = orgs[0], orgs[1]

from docx import Document
doc = Document()
for num, date, item, price in (("МУЛЬТИ-А/1", "01.07.2026", "Товар А", 100), ("МУЛЬТИ-Б/2", "05.07.2026", "Товар Б", 200)):
    doc.add_paragraph(f"КОНТРАКТ ПОСТАВКИ № {num}")
    doc.add_paragraph(f"«{date[:2]}» июля 2026 г. Заказчик ИНН {customer_inn}, Поставщик ИНН {supplier_inn}.")
    doc.add_paragraph(f"Спецификация: {item} — 2 шт по {price},00 руб. Срок поставки: до 30.07.2026.")
buf = io.BytesIO()
doc.save(buf)
content = buf.getvalue()
PENDING.mkdir(parents=True, exist_ok=True)
(PENDING / f"{DOC_ID}.bin").write_bytes(content)
(PENDING / f"{DOC_ID}.json").write_text(json.dumps({
    "id": DOC_ID, "filename": "ТЕСТ мульти — УДАЛИТЬ.docx", "size_bytes": len(content),
    "uploaded_at": int(time.time()),
    "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}, ensure_ascii=False), encoding="utf-8")

ids = []
try:
    def extracted(num, item, price):
        return {
            "contractNumber": num, "contractDate": "01.07.2026", "deadlineDate": "30.07.2026",
            "customerInn": customer_inn, "supplierInn": supplier_inn, "vatPercent": "0",
            "items": [{"name": item, "docName": item, "unit": "шт", "qty": 2, "price": price}],
        }

    # contract A — keepDocument=true
    res_a = call("save_contract_from_incoming_document", {
        "documentId": DOC_ID, "keepDocument": True, "extracted": extracted("МУЛЬТИ-А/1", "Товар А", 100),
    })
    ids.append(str(res_a["contract"]["id"]))
    assert res_a["status"] == "created" and res_a["documentRemovedFromIncoming"] is False, res_a["status"]
    docs = call("list_incoming_contract_documents", {})
    mine = next((d for d in docs["documents"] if d["id"] == DOC_ID), None)
    assert mine and mine["processed"] is False, f"doc must stay unprocessed in incoming: {mine}"
    files_a = call("get_contract_files", {"contractId": ids[0]})
    assert files_a["count"] >= 1, files_a
    print(f"contract A OK: doc stayed in incoming (processed={mine['processed']}), file attached to A")

    # contract B — final save, no keepDocument
    res_b = call("save_contract_from_incoming_document", {
        "documentId": DOC_ID, "extracted": extracted("МУЛЬТИ-Б/2", "Товар Б", 200),
    })
    ids.append(str(res_b["contract"]["id"]))
    assert res_b["status"] == "created" and res_b["documentRemovedFromIncoming"] is True
    docs2 = call("list_incoming_contract_documents", {})
    assert not any(d["id"] == DOC_ID for d in docs2["documents"]), "doc must be consumed after final save"
    files_b = call("get_contract_files", {"contractId": ids[1]})
    assert files_b["count"] >= 1, files_b
    assert res_a["deadlineSaved"] and res_b["deadlineSaved"], (res_a.get("deadlineSaved"), res_b.get("deadlineSaved"))
    print("contract B OK: doc consumed after final save, file attached to B")
    print(f"A total={res_a['totalAmount']}, B total={res_b['totalAmount']}, deadlines={res_a['deadlineSaved']}/{res_b['deadlineSaved']}")
    print("MULTI_CONTRACT_OK")
finally:
    import psycopg
    db_url = (env.get("DATABASE_URL") or "").replace("postgresql+psycopg2://", "postgresql://")
    with psycopg.connect(db_url) as conn:
        with conn.cursor() as cur:
            for cid in ids:
                cur.execute("SELECT storage_key FROM contract_files WHERE contract_id = %s", (cid,))
                for (key,) in cur.fetchall():
                    p = (APP / key).resolve()
                    if "contract_files_store" in str(p) and p.exists():
                        p.unlink()
                cur.execute("DELETE FROM inventory_movements WHERE contract_id = %s", (cid,))
                cur.execute("DELETE FROM delivery_stage_items WHERE stage_id IN (SELECT id FROM delivery_stages WHERE contract_id = %s)", (cid,))
                cur.execute("DELETE FROM delivery_stages WHERE contract_id = %s", (cid,))
                cur.execute("DELETE FROM contract_items WHERE contract_id = %s", (cid,))
                cur.execute("DELETE FROM contract_files WHERE contract_id = %s", (cid,))
                cur.execute("DELETE FROM contracts WHERE id = %s", (cid,))
            cur.execute("DELETE FROM pending_document_ai_state WHERE doc_id = %s", (DOC_ID,))
        conn.commit()
    for suffix in (".json", ".bin", ".preview.bin"):
        p = PENDING / f"{DOC_ID}{suffix}"
        if p.exists():
            p.unlink()
    print("cleanup done")
