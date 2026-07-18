"""Live e2e: fast contract intake + real org linking + view-as-images (runs on miramed32).

Seeds a synthetic contract DOCX as a pending document using REAL directory org INNs
(no new orgs created, no DaData calls), then:
  view (images) -> one-shot save_contract_from_incoming_document -> asserts real org
  cards linked (not placeholder), doc marked processed, file attached -> full SQL cleanup.
Also runs view_incoming_contract_document on a real scanned/PDF incoming doc (read-only).
"""
import base64
import io
import json
import re
import time
import urllib.request
from pathlib import Path

APP = Path("/var/www/prostye-postavki/app")
import sys
sys.path.insert(0, str(APP))

PENDING = APP / "backend" / "data" / "pending_contract_files"
DOC_ID = "smoketest-intake-fast"

env = {}
for line in (APP / ".env.local").read_text(encoding="utf-8", errors="replace").splitlines():
    if "=" in line and not line.startswith("#"):
        k, _, v = line.partition("=")
        env[k.strip()] = v.strip().strip('"')
secret = env["MCP_SERVER_SECRET"]
URL = f"http://127.0.0.1:8000/mcp/{secret}"


def call_raw(name, arguments):
    body = json.dumps({"jsonrpc": "2.0", "id": 1, "method": "tools/call",
                       "params": {"name": name, "arguments": arguments}}).encode()
    req = urllib.request.Request(URL, data=body, headers={"content-type": "application/json"})
    with urllib.request.urlopen(req, timeout=180) as resp:
        payload = json.loads(resp.read().decode())
    if "error" in payload:
        raise RuntimeError(f"{name}: {payload['error']}")
    return payload["result"]


def call(name, arguments):
    result = call_raw(name, arguments)
    if isinstance(result, dict) and result.get("isError"):
        raise RuntimeError(f"{name}: {result}")
    return json.loads(result["content"][0]["text"])


# --- pick two REAL active orgs with INN from the directory (read-only) ---
rows = call("read_table_rows", {"table": "organizations", "limit": 200})
orgs = []
for r in rows.get("rows", []):
    inn = re.sub(r"\D", "", str(r.get("inn") or ""))
    name = str(r.get("name") or "")
    if len(inn) in (10, 12) and not name.startswith("__SYSTEM_") and r.get("is_active") in (True, None):
        orgs.append({"inn": inn, "name": name})
assert len(orgs) >= 2, f"need 2 real orgs, got {len(orgs)}"
customer, supplier = orgs[0], orgs[1]
print(f"customer org: {customer['name'][:50]!r} ИНН {customer['inn']}")
print(f"supplier org: {supplier['name'][:50]!r} ИНН {supplier['inn']}")

# --- seed synthetic contract DOCX with those INNs ---
from docx import Document

doc = Document()
doc.add_paragraph("ДОГОВОР ПОСТАВКИ № ТЕСТ-ИНТЕЙК/1")
doc.add_paragraph("г. Брянск, «10» июля 2026 г.")
doc.add_paragraph(
    f"Заказчик: {customer['name']}, ИНН {customer['inn']}. "
    f"Поставщик: {supplier['name']}, ИНН {supplier['inn']}."
)
doc.add_paragraph("Спецификация: Тестовая позиция интейка — 3 шт по 250,00 руб. Итого 750,00 руб.")
doc.add_paragraph("Срок поставки: до 25.07.2026.")
buf = io.BytesIO()
doc.save(buf)
content = buf.getvalue()
PENDING.mkdir(parents=True, exist_ok=True)
(PENDING / f"{DOC_ID}.bin").write_bytes(content)
(PENDING / f"{DOC_ID}.json").write_text(json.dumps({
    "id": DOC_ID, "filename": "ТЕСТ интейк — УДАЛИТЬ.docx",
    "size_bytes": len(content), "uploaded_at": int(time.time()),
    "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}, ensure_ascii=False), encoding="utf-8")

created_contract_id = ""
try:
    # --- view seeded DOCX as images (office -> pdf -> jpeg pipeline) ---
    view = call_raw("view_incoming_contract_document", {"documentId": DOC_ID, "pages": [1]})
    blocks = view["content"]
    images = [b for b in blocks if b.get("type") == "image"]
    assert images, f"no image blocks: {[b.get('type') for b in blocks]}"
    jpeg = base64.b64decode(images[0]["data"])
    assert jpeg[:2] == b"\xff\xd8", "not a JPEG"
    print(f"view (docx->pdf->jpeg) OK: {len(images)} page image(s), first {len(jpeg)} bytes")

    # --- one-shot intake ---
    res = call("save_contract_from_incoming_document", {
        "documentId": DOC_ID,
        "extracted": {
            "contractNumber": "ТЕСТ-ИНТЕЙК/1",
            "contractDate": "10.07.2026",
            "deadlineDate": "25.07.2026",
            "customerInn": customer["inn"],
            "supplierInn": supplier["inn"],
            "vatPercent": "0",
            "items": [{"name": "Тестовая позиция интейка", "docName": "Тестовая позиция интейка",
                       "unit": "шт", "qty": 3, "price": 250}],
        },
    })
    assert res["status"] == "created", res["status"]
    created_contract_id = str(res["contract"]["id"])
    lo = res["linkedOrganizations"]
    print(f"contract created: {created_contract_id}")
    print(f"  linked customer: {lo['customer']['name'][:50]!r} inn={lo['customer']['inn']} placeholder={lo['customer']['isPlaceholder']}")
    print(f"  linked supplier: {lo['supplier']['name'][:50]!r} inn={lo['supplier']['inn']} placeholder={lo['supplier']['isPlaceholder']}")
    assert lo["customer"]["isPlaceholder"] is False and re.sub(r"\D", "", lo["customer"]["inn"]) == customer["inn"], lo
    assert lo["supplier"]["isPlaceholder"] is False and re.sub(r"\D", "", lo["supplier"]["inn"]) == supplier["inn"], lo
    assert res["itemsSaved"] and res["itemsSaved"][0]["qty"] == 3 and res["itemsSaved"][0]["price"] == 250, res["itemsSaved"]
    assert res["totalAmount"] == 750.0, res["totalAmount"]
    assert res["sumWarnings"] == [], res["sumWarnings"]
    print(f"  itemsSaved OK: {res['itemsSaved']}, totalAmount={res['totalAmount']}, sumWarnings=[]")

    assert res["documentRemovedFromIncoming"] is True
    # --- doc consumed: gone from incoming (no ghost meta), file lives in the contract ---
    docs = call("list_incoming_contract_documents", {"limit": 200})
    assert not any(d["id"] == DOC_ID for d in docs["documents"]), "ghost meta re-created"
    import os
    assert not os.path.exists(str(PENDING / f"{DOC_ID}.json")), "ghost meta file exists"
    print(f"doc consumed cleanly, no ghost (list count={docs['count']}, unprocessed={docs['unprocessedCount']})")
    for d in docs["documents"]:
        print(f"  real incoming doc visible: {d['filename']!r} processed={d['processed']}")

    # --- OVERWRITE by contractId only (documentId consumed) ---
    res2 = call("save_contract_from_incoming_document", {
        "contractId": created_contract_id,
        "extracted": {
            "contractNumber": "ТЕСТ-ИНТЕЙК/1",
            "contractDate": "10.07.2026",
            "deadlineDate": "28.07.2026",
            "customerInn": customer["inn"],
            "supplierInn": supplier["inn"],
            "vatPercent": "0",
            "items": [{"name": "Тестовая позиция интейка", "docName": "Тестовая позиция интейка",
                       "unit": "шт", "qty": 5, "price": 200}],
        },
    })
    assert res2["status"] == "updated", res2["status"]
    assert res2["itemsSaved"][0]["qty"] == 5 and res2["totalAmount"] == 1000.0, res2["itemsSaved"]
    assert res2["linkedOrganizations"]["customer"]["isPlaceholder"] is False
    files2 = call("get_contract_files", {"contractId": created_contract_id})
    assert files2["count"] >= 1, "attached file lost on overwrite"
    print(f"OVERWRITE by contractId OK: qty->5, total->1000.0, file kept ({files2['count']})")

    # --- file attached to the contract ---
    files = call("get_contract_files", {"contractId": created_contract_id})
    assert files["count"] >= 1 and any(f["isDocx"] for f in files["files"]), files
    print(f"contract file attached OK ({files['files'][0]['filename']!r})")

    # --- contract visible via get_contracts ---
    got = call("get_contracts", {"query": "ТЕСТ-ИНТЕЙК", "limit": 10})
    print(f"get_contracts finds it: {bool(got.get('contracts') or got.get('rows') or got.get('count'))}")

    # --- view a REAL scanned/PDF incoming document (read-only) ---
    all_docs = call("list_incoming_contract_documents", {"includeProcessed": True, "limit": 100})
    real_pdf = next((d for d in all_docs["documents"]
                     if d["id"] != DOC_ID and (d.get("previewType") == "pdf" or str(d.get("mimeType", "")).endswith("pdf"))), None)
    if real_pdf:
        view2 = call_raw("view_incoming_contract_document", {"documentId": real_pdf["id"], "maxPages": 2})
        imgs2 = [b for b in view2["content"] if b.get("type") == "image"]
        assert imgs2 and base64.b64decode(imgs2[0]["data"])[:2] == b"\xff\xd8"
        print(f"view on REAL doc {real_pdf['filename']!r} OK: {len(imgs2)} page image(s)")
    else:
        print("no real PDF/scan incoming docs to view — skipped")

    print("LIVE_INTAKE_OK")
finally:
    # --- full cleanup: contract + files + seeded pending doc ---
    import traceback
    traceback.print_exc()
    import psycopg
    db_url = (env.get("DATABASE_URL") or env.get("DB_URL") or "").replace("postgresql+psycopg2://", "postgresql://").replace("postgresql+psycopg://", "postgresql://")
    with psycopg.connect(db_url) as conn:
        if created_contract_id:
            with conn.cursor() as cur:
                cur.execute("SELECT storage_key FROM contract_files WHERE contract_id = %s", (created_contract_id,))
                for (key,) in cur.fetchall():
                    p = (APP / key).resolve()
                    if str(p).startswith(str((APP / "backend" / "data" / "contract_files_store").resolve())) and p.exists():
                        p.unlink()
                cur.execute("DELETE FROM inventory_movements WHERE contract_id = %s", (created_contract_id,))
                cur.execute("DELETE FROM delivery_stage_items WHERE stage_id IN (SELECT id FROM delivery_stages WHERE contract_id = %s)", (created_contract_id,))
                cur.execute("DELETE FROM delivery_stages WHERE contract_id = %s", (created_contract_id,))
                cur.execute("DELETE FROM contract_items WHERE contract_id = %s", (created_contract_id,))
                cur.execute("DELETE FROM contract_files WHERE contract_id = %s", (created_contract_id,))
                cur.execute("DELETE FROM contracts WHERE id = %s", (created_contract_id,))
        with conn.cursor() as cur:
            cur.execute("DELETE FROM pending_document_ai_state WHERE doc_id = %s", (DOC_ID,))
        conn.commit()
    for suffix in (".json", ".bin", ".preview.bin"):
        p = PENDING / f"{DOC_ID}{suffix}"
        if p.exists():
            p.unlink()
    print("cleanup done (contract, files, seeded doc removed)")
